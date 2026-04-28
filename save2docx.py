# -*- coding: utf-8 -*-

import json
import math
import os
import re
import glob
from datetime import datetime

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.shared import qn
from docx.shared import Inches, Pt, Cm, RGBColor

CHINESE_FONT = '仿宋'
ENGLISH_FONT = 'Times New Roman'

# =========================================================
# 基础工具
# =========================================================

def _ensure_json_suffix(filename: str) -> str:
    """确保文件名以 .json 结尾"""
    return filename if filename.endswith('.json') else f'{filename}.json'


def save_state(state, state_path, filename='state'):
    """
    保存程序状态到指定文件。
    """
    os.makedirs(state_path, exist_ok=True)
    file_path = os.path.join(state_path, _ensure_json_suffix(filename))
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(state, f, indent=4, ensure_ascii=False)


def load_state(state_path, filename='state'):
    """
    从指定目录加载程序状态。
    """
    file_path = os.path.join(state_path, _ensure_json_suffix(filename))
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None


def find_json_path(candidate_dirs, filename='state'):
    """
    在多个候选目录中查找 json 文件，返回第一个命中的完整路径。
    """
    filename = _ensure_json_suffix(filename)
    for d in candidate_dirs:
        if not d:
            continue
        file_path = os.path.join(d, filename)
        if os.path.exists(file_path):
            return file_path
    return None


def load_json_from_candidates(candidate_dirs, filename='state'):
    """
    从多个候选目录加载 json，并返回 (data, file_path)。
    """
    file_path = find_json_path(candidate_dirs, filename)
    if file_path is None:
        return None, None

    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f), file_path


def is_null_like(value):
    """判断 None / NaN"""
    if value is None:
        return True
    if isinstance(value, float) and math.isnan(value):
        return True
    return False


def _safe_int(value):
    try:
        return int(value)
    except Exception:
        return None


def fmt_num(value, digits=2):
    """格式化数值"""
    if is_null_like(value):
        return "-"
    if isinstance(value, bool):
        return "True" if value else "False"
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    try:
        v = float(value)
        if math.isnan(v):
            return "-"
        return f"{v:.{digits}f}"
    except Exception:
        return str(value)


def format_value_for_table(value, digits=4):
    """把任意值格式化成适合 Word 表格显示的字符串"""
    if value is None:
        return "-"
    if isinstance(value, bool):
        return "True" if value else "False"
    if isinstance(value, int) and not isinstance(value, bool):
        return str(value)
    if isinstance(value, float):
        return fmt_num(value, digits)
    if isinstance(value, (dict, list, tuple, set)):
        try:
            if isinstance(value, set):
                value = list(value)
            return json.dumps(value, ensure_ascii=False)
        except Exception:
            return str(value)
    return str(value)


def format_config_value(value):
    """把配置值格式化成适合 Word 表格显示的字符串"""
    if value is None:
        return "-"
    if isinstance(value, bool):
        return "True" if value else "False"
    if isinstance(value, (dict, list, tuple, set)):
        try:
            if isinstance(value, set):
                value = list(value)
            return json.dumps(value, ensure_ascii=False)
        except Exception:
            return str(value)
    return str(value)


def flatten_params_dict(d, parent_key="", sep="."):
    """
    递归展开嵌套字典，把每个叶子参数展开成一行。
    例如：
    customize.jv_params.test_periods_start -> 100
    """
    items = []
    if not isinstance(d, dict):
        items.append((parent_key if parent_key else "value", d))
        return items

    for k, v in d.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else str(k)
        if isinstance(v, dict):
            items.extend(flatten_params_dict(v, new_key, sep=sep))
        else:
            items.append((new_key, v))
    return items


def extract_source_index(source_name):
    """从 '1_4FGL_xxx' 里提取 1；失败则返回 None"""
    if not source_name:
        return None
    m = re.match(r'^(\d+)_', str(source_name))
    if not m:
        return None
    try:
        return int(m.group(1))
    except Exception:
        return None


def format_date_like(value):
    """
    将日期值格式化为 YYYY-MM-DD。
    支持：
    - [2026, 3, 31]
    - (2026, 3, 31)
    - "2026,3,31"
    - "2026-3-31"
    - "2026/3/31"
    - datetime
    """
    if is_null_like(value):
        return "-"

    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")

    if isinstance(value, (list, tuple)):
        if len(value) >= 3:
            y = _safe_int(value[0])
            m = _safe_int(value[1])
            d = _safe_int(value[2])
            if y is not None and m is not None and d is not None:
                try:
                    return datetime(y, m, d).strftime("%Y-%m-%d")
                except Exception:
                    pass
        return format_value_for_table(value)

    if isinstance(value, str):
        s = value.strip()
        if not s:
            return "-"
        # 仅针对明显的日期格式做转换
        patterns = [
            r'^\d{4},\d{1,2},\d{1,2}$',
            r'^\d{4}-\d{1,2}-\d{1,2}$',
            r'^\d{4}/\d{1,2}/\d{1,2}$',
        ]
        if any(re.match(p, s) for p in patterns):
            parts = re.split(r'[,/-]', s)
            if len(parts) >= 3:
                y = _safe_int(parts[0])
                m = _safe_int(parts[1])
                d = _safe_int(parts[2])
                if y is not None and m is not None and d is not None:
                    try:
                        return datetime(y, m, d).strftime("%Y-%m-%d")
                    except Exception:
                        pass
        try:
            # 兼容 ISO 格式
            return datetime.fromisoformat(s).strftime("%Y-%m-%d")
        except Exception:
            return s

    return str(value)


def format_date_range(start_value, end_value):
    """格式化日期范围"""
    start_text = format_date_like(start_value)
    end_text = format_date_like(end_value)

    if start_text == "-" and end_text == "-":
        return "-"
    if start_text != "-" and end_text != "-":
        return f"{start_text} 至 {end_text}"
    if start_text != "-":
        return f"{start_text} 至 -"
    return f"- 至 {end_text}"


def get_source_override_global_dates(config, source_name):
    """
    从 config.source_overrides 中提取某个源的 global.start_date / end_date。
    """
    if not isinstance(config, dict):
        return None, None, None

    idx = extract_source_index(source_name)
    if idx is None:
        return None, None, None

    source_overrides = config.get("source_overrides", {}) or {}
    override = source_overrides.get(str(idx), {}) or {}
    global_override = override.get("global", {}) or {}

    return (
        global_override.get("start_date"),
        global_override.get("end_date"),
        global_override,
    )


def resolve_source_date_info(source_name, source_result, config):
    """
    计算单个源的日期信息：
    - global range
    - config override range
    - applied range (state.json 里的实际应用范围)
    - effective range (优先 applied，其次 override，其次 global)
    """
    source_result = source_result or {}
    global_cfg = (config or {}).get("global", {}) if isinstance(config, dict) else {}

    global_start = global_cfg.get("start_date")
    global_end = global_cfg.get("end_date")

    override_start, override_end, override_obj = get_source_override_global_dates(config, source_name)

    applied_start = source_result.get("applied_start_date")
    applied_end = source_result.get("applied_end_date")

    effective_start = applied_start if not is_null_like(applied_start) else (
        override_start if not is_null_like(override_start) else global_start
    )
    effective_end = applied_end if not is_null_like(applied_end) else (
        override_end if not is_null_like(override_end) else global_end
    )

    return {
        "global_start": global_start,
        "global_end": global_end,
        "override_start": override_start,
        "override_end": override_end,
        "override_obj": override_obj,
        "applied_start": applied_start,
        "applied_end": applied_end,
        "effective_start": effective_start,
        "effective_end": effective_end,
        "global_range": format_date_range(global_start, global_end),
        "override_range": format_date_range(override_start, override_end),
        "applied_range": format_date_range(applied_start, applied_end),
        "effective_range": format_date_range(effective_start, effective_end),
    }


def find_processed_file_for_source(source_name, processed_files):
    """
    从 processed_files 中尽量匹配出与 source_name 对应的源文件名。
    """
    if not processed_files:
        return None

    variants = set()
    if source_name:
        variants.add(str(source_name))
        short_name = re.sub(r'^\d+_', '', str(source_name))
        variants.add(short_name)

        if "_" in short_name:
            parts = short_name.split("_")
            if len(parts) >= 2:
                variants.add("_".join(parts[1:]))
                variants.add(parts[-1])

    candidates = []
    for idx, file_name in enumerate(processed_files):
        stem = os.path.splitext(os.path.basename(file_name))[0].lower()
        score = 0
        if any(v.lower() in stem for v in variants):
            score += 80
        if source_name and str(source_name).lower() in stem:
            score += 40
        short_name = re.sub(r'^\d+_', '', str(source_name)) if source_name else ""
        if short_name and short_name.lower() in stem:
            score += 20

        if score > 0:
            candidates.append((score, -idx, file_name))

    if not candidates:
        return None

    candidates.sort(key=lambda x: (-x[0], x[1]))
    return candidates[0][2]


def get_ordered_result_sources(state_dicts):
    """按 results 的插入顺序获取成功源列表。"""
    results = state_dicts.get("results", {}) or {}
    if isinstance(results, dict):
        return list(results.keys())
    return []


def ensure_list(obj):
    """确保返回 list"""
    if obj is None:
        return []
    if isinstance(obj, list):
        return obj
    if isinstance(obj, tuple):
        return list(obj)
    return [obj]


# =========================================================
# 字体与段落工具
# =========================================================

def set_run_font(run, chinese_font=CHINESE_FONT, english_font=ENGLISH_FONT,
                 size=None, bold=None, italic=None, color=None):
    """
    同时设置中文字体和英文字体：
    - 中文：eastAsia
    - 英文：ascii / hAnsi / cs

    这样中英文混排时会自动分别使用对应字体。
    """
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)

    rFonts.set(qn('w:eastAsia'), chinese_font)
    rFonts.set(qn('w:ascii'), english_font)
    rFonts.set(qn('w:hAnsi'), english_font)
    rFonts.set(qn('w:cs'), english_font)

    run.font.name = english_font

    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_text(paragraph, text, chinese_font=CHINESE_FONT, english_font=ENGLISH_FONT,
                       size=Pt(10), bold=False, italic=False, color=None,
                       align=WD_ALIGN_PARAGRAPH.LEFT, clear=True):
    """
    给段落写入文本并设置字体。
    """
    if clear:
        try:
            paragraph.clear()
        except Exception:
            pass

    run = paragraph.add_run(str(text))
    set_run_font(
        run,
        chinese_font=chinese_font,
        english_font=english_font,
        size=size,
        bold=bold,
        italic=italic,
        color=color
    )
    paragraph.alignment = align
    return run


def add_styled_heading(doc, text, level=1, center=False,
                       chinese_font=CHINESE_FONT, english_font=ENGLISH_FONT,
                       size=None, bold=True):
    """
    添加统一字体的标题。
    """
    if size is None:
        size = {1: Pt(14), 2: Pt(12), 3: Pt(11)}.get(level, Pt(10))

    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(str(text))
    set_run_font(
        run,
        chinese_font=chinese_font,
        english_font=english_font,
        size=size,
        bold=bold
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    return paragraph


def set_cell_text(cell, text, chinese_font=CHINESE_FONT, english_font=ENGLISH_FONT,
                  size=Pt(10), bold=False, italic=False, color=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    """
    给单元格写入文本并统一字体。
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]
    return set_paragraph_text(
        paragraph, text,
        chinese_font=chinese_font,
        english_font=english_font,
        size=size,
        bold=bold,
        italic=italic,
        color=color,
        align=align,
        clear=False
    )


# =========================================================
# 占位图片
# =========================================================

def create_placeholder_image(text="No Image", width=520, height=360, temp_dir='temp_images'):
    """创建默认占位图片"""
    img = Image.new('RGB', (width, height), color='white')
    d = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        font = ImageFont.load_default()

    try:
        bbox = d.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
    except Exception:
        text_width = d.textlength(text, font=font)
        text_height = 20

    d.text(
        ((width - text_width) / 2, (height - text_height) / 2),
        text,
        fill='black',
        font=font
    )

    os.makedirs(temp_dir, exist_ok=True)
    safe_name = re.sub(r'[^A-Za-z0-9_\-]+', '_', text)
    path = os.path.join(temp_dir, f'placeholder_{safe_name}.png')
    img.save(path)
    return path


# =========================================================
# Word 格式工具
# =========================================================

def set_no_border_table(table):
    """设置表格为无边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr.xpath('w:tblBorders') == []:
        tblBorders = parse_xml(
            r'<w:tblBorders {}>'
            r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            r'</w:tblBorders>'.format(nsdecls('w'))
        )
        tblPr.append(tblBorders)
    else:
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            try:
                border_tag = tblPr.xpath('w:tblBorders/w:{}'.format(border))[0]
                border_tag.set(qn('w:val'), 'none')
                border_tag.set(qn('w:sz'), '0')
                border_tag.set(qn('w:space'), '0')
                border_tag.set(qn('w:color'), 'auto')
            except Exception:
                pass


def add_bookmark(paragraph, bookmark_name, bookmark_id=0):
    """在段落中添加书签"""
    run = paragraph.add_run()
    tag = run._r
    start = parse_xml(
        r'<w:bookmarkStart {} w:id="{}" w:name="{}"/>'.format(nsdecls('w'), bookmark_id, bookmark_name)
    )
    end = parse_xml(r'<w:bookmarkEnd {} w:id="{}"/>'.format(nsdecls('w'), bookmark_id))
    tag.append(start)
    tag.append(end)


def setup_footer(doc):
    """设置页脚和页码"""
    section = doc.sections[0]
    footer = section.footer

    for paragraph in footer.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    paragraph = footer.paragraphs[0] if len(footer.paragraphs) > 0 else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    fld_char1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    fld_char2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))
    fld_char3 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))

    run._r.append(fld_char1)
    run._r.append(parse_xml(r'<w:instrText {} xml:space="preserve">PAGE </w:instrText>'.format(nsdecls('w'))))
    run._r.append(fld_char2)
    run._r.append(parse_xml(r'<w:t {}>1</w:t>'.format(nsdecls('w'))))
    run._r.append(fld_char3)


def setup_header(doc, program_name="Program"):
    """设置页眉，左边写程序名，右边写日期"""
    section = doc.sections[0]
    header = section.header

    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    table = header.add_table(rows=1, cols=2, width=Inches(7.4))
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)
    set_no_border_table(table)

    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    left_cell = table.cell(0, 0)
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_paragraph.add_run(program_name)
    set_run_font(left_run, size=Pt(10), bold=True)

    right_cell = table.cell(0, 1)
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_paragraph.add_run(current_date)
    set_run_font(right_run, size=Pt(10), bold=False)


def set_cell_border(cell, border_dict):
    """
    Set cell's border according to border_dict

    border_dict: {
        'top': {'sz': 8, 'val': 'single', 'color': '#000000'},
        'bottom': {'sz': 8, 'val': 'single', 'color': '#000000'},
        'start': {'sz': 0, 'val': 'none', 'color': '#FFFFFF'},
        'end': {'sz': 0, 'val': 'none', 'color': '#FFFFFF'}
    }
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = border_dict.get(edge)
        if edge_data:
            tag = f'w:{edge}'

            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key, value in edge_data.items():
                element.set(qn(f'w:{key}'), value)


# =========================================================
# 配置表：一参数一行
# =========================================================

def add_parameters_table(doc, params_dict, title="Parameter Settings"):
    """
    将参数字典添加为表格到 Word 文档。

    改进点：
    - 不再把嵌套字典直接显示成一个单元格
    - 每个 leaf 参数单独占一行
    - 例如：
      customize.jv_params.test_periods_start -> 100
    """
    add_styled_heading(doc, title, level=1, center=False, size=Pt(14), bold=True)

    if not isinstance(params_dict, dict) or not params_dict:
        p = doc.add_paragraph()
        r = p.add_run("No parameter settings available.")
        set_run_font(r, chinese_font=CHINESE_FONT, english_font=ENGLISH_FONT)

        doc.add_page_break()
        return

    for section_name, section_params in params_dict.items():
        add_styled_heading(doc, section_name, level=2, center=False, size=Pt(12), bold=True)

        flat_items = flatten_params_dict(section_params, parent_key="")

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        set_cell_text(hdr_cells[0], 'Parameter', size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(hdr_cells[1], 'Value', size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        for param_name, param_value in flat_items:
            row_cells = table.add_row().cells
            set_cell_text(row_cells[0], param_name, size=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(row_cells[1], format_config_value(param_value), size=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

        doc.add_paragraph()

    doc.add_page_break()


# =========================================================
# 图片与结果辅助
# =========================================================

def _coerce_float(value):
    """尽量把值转成 float，失败就返回 None。"""
    if is_null_like(value):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        s = value.strip()
        try:
            return float(s)
        except Exception:
            m = re.search(r'[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?', s)
            if m:
                try:
                    return float(m.group(0))
                except Exception:
                    return None
    return None


def extract_beta_fields(container):
    """
    从结果里提取 beta 和 beta error。

    兼容：
    - state_result["Beta"] = {"beta_best": ..., "beta_err": ...}
    - 未来可能出现的 beta / beta_error 命名
    - 如果直接传入 source_result，也会自动从 Beta 或 LSP 里找
    """
    if not isinstance(container, dict):
        return None, None

    beta = None
    beta_err = None

    # 1) 如果直接传的是 source_result，优先看 Beta 子块
    sub = container.get("Beta")
    if isinstance(sub, dict):
        beta = sub.get("beta_best", sub.get("beta"))
        beta_err = sub.get("beta_err", sub.get("beta_error"))

    # 2) 如果传的是 Beta 子块本身
    if beta is None:
        beta = container.get("beta_best", container.get("beta"))
    if beta_err is None:
        beta_err = container.get("beta_err", container.get("beta_error"))

    # 3) 兼容旧结构：LSP/PSRESP/Log 中可能存在 beta
    if beta is None or beta_err is None:
        for k in ("LSP", "lsp", "PSRESP", "psresp", "Log", "log"):
            sub2 = container.get(k)
            if isinstance(sub2, dict):
                if beta is None:
                    beta = sub2.get("beta_best", sub2.get("beta"))
                if beta_err is None:
                    beta_err = sub2.get("beta_err", sub2.get("beta_error"))

    return _coerce_float(beta), _coerce_float(beta_err)


def format_beta_text(beta, beta_err):
    """把 beta 和误差格式化成一行文本。"""
    if is_null_like(beta) and is_null_like(beta_err):
        return "-"
    if not is_null_like(beta) and not is_null_like(beta_err):
        return f"{fmt_num(beta, 4)} ± {fmt_num(beta_err, 4)}"
    if not is_null_like(beta):
        return fmt_num(beta, 4)
    if not is_null_like(beta_err):
        return f"± {fmt_num(beta_err, 4)}"
    return "-"


def summarize_method_result(method_name, method_result):
    """
    把某个方法的结果压缩成适合 Overview 展示的一行摘要。
    """
    if method_result is None:
        if method_name == "Light_Plot":
            return "光变曲线"
        return "-"

    if method_name == "Light_Plot":
        return "光变曲线"

    if method_name == "Beta":
        beta, beta_err = extract_beta_fields(method_result)
        return f"β {format_beta_text(beta, beta_err)}"

    if method_name == "LSP":
        periods = method_result.get("periods") or []
        beta, beta_err = extract_beta_fields(method_result)
        beta_text = ""
        if not is_null_like(beta) or not is_null_like(beta_err):
            beta_text = f", β {format_beta_text(beta, beta_err)}"

        if not periods:
            return f"未检出显著周期{beta_text}" if beta_text else "未检出显著周期"

        first = periods[0]
        p = first.get("period")
        return f"{len(periods)}个周期，首个周期 {fmt_num(p)} d{beta_text}"

    if method_name == "Jurkevich":
        if not isinstance(method_result, dict):
            return "无结果"
        p = method_result.get("period")
        if is_null_like(p):
            return "无结果"
        return f"1个周期，首个周期 {fmt_num(p)} d"

    if method_name == "DCF":
        if not isinstance(method_result, dict):
            return "无结果"
        periods = method_result.get("period") or []
        if not periods:
            return "未检出显著周期"
        first = periods[0]
        p = first.get("period")
        return f"{len(periods)}个周期，首个周期 {fmt_num(p)} d"

    if method_name == "WWZ":
        if not isinstance(method_result, dict):
            return "无结果"
        peaks = method_result.get("result") or []
        if not peaks:
            return "未检出显著周期"
        first = peaks[0]
        p = first.get("period")
        return f"{len(peaks)}个周期，首个周期 {fmt_num(p)} d"

    return "无结果"


PLOT_SPECS = [
    {
        "label": "Light_Plot",
        "folder": "Light_Plot",
        "result_key": None,
        "summary_method": "Light_Plot",
        "must_have": None,
        "must_not_have": ["psresp", "log", "beta", "slope"],
    },
    {
        "label": "LSP",
        "folder": "LSP",
        "result_key": "LSP",
        "summary_method": "LSP",
        "must_have": None,
        "must_not_have": ["psresp", "log", "beta", "slope"],
    },
    {
        "label": "β斜率",
        "folder": "LSP",
        "result_key": "Beta",
        "summary_method": "Beta",
        "must_have": ["psresp", "log"],
        "must_not_have": [],
    },
    {
        "label": "Jurkevich",
        "folder": "Jurkevich",
        "result_key": "Jurkevich",
        "summary_method": "Jurkevich",
        "must_have": None,
        "must_not_have": None,
    },
    {
        "label": "DCF",
        "folder": "DCF",
        "result_key": "DCF",
        "summary_method": "DCF",
        "must_have": None,
        "must_not_have": None,
    },
    {
        "label": "WWZ",
        "folder": "WWZ",
        "result_key": "WWZ",
        "summary_method": "WWZ",
        "must_have": None,
        "must_not_have": None,
    },
]


def get_source_variants(source_name):
    """
    生成多个可能的源名匹配串，尽量兼容不同图片命名方式。
    例如：
    - 1_4FGL_J0003.3-1928
    - 4FGL_J0003.3-1928
    - J0003.3-1928
    """
    variants = set()
    if not source_name:
        return variants

    variants.add(source_name)
    short_name = re.sub(r'^\d+_', '', source_name)
    variants.add(short_name)

    if "_" in short_name:
        parts = short_name.split("_")
        if len(parts) >= 2:
            variants.add("_".join(parts[1:]))
            variants.add(parts[-1])

    return {v for v in variants if v}


def find_best_image_in_folder(data_path, folder_name, source_name,
                              must_have=None, must_not_have=None):
    """
    只在指定 folder_name 里找图，避免跨目录误匹配。

    规则：
    - 必须包含 source_name 相关片段
    - must_have：如果指定，则文件名必须至少命中一个关键词
    - must_not_have：如果指定，则文件名命中这些关键词会被排除
    """
    search_dirs = []

    folder_path = os.path.join(data_path, folder_name)
    if os.path.isdir(folder_path):
        search_dirs.append(folder_path)
    elif os.path.isdir(data_path):
        # 兜底：如果对应 folder 不存在，才退回 data_path
        search_dirs.append(data_path)

    source_variants = list(get_source_variants(source_name))
    must_have = [s.lower() for s in (must_have or [])]
    must_not_have = [s.lower() for s in (must_not_have or [])]

    image_exts = (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff")
    candidates = []

    seen = set()
    for sd in search_dirs:
        for path in glob.glob(os.path.join(sd, "**", "*"), recursive=True):
            if not os.path.isfile(path):
                continue
            if os.path.splitext(path)[1].lower() not in image_exts:
                continue
            if path in seen:
                continue
            seen.add(path)

            filename = os.path.basename(path)
            stem = os.path.splitext(filename)[0].lower()

            # 1) 必须能匹配到源名
            if not any(v.lower() in stem for v in source_variants):
                continue

            # 2) 排除不想要的后缀
            if must_not_have and any(term in stem for term in must_not_have):
                continue

            # 3) 如果指定了 must_have，则至少命中一个关键词
            if must_have and not any(term in stem for term in must_have):
                continue

            score = 0

            # 源名匹配最重要
            score += 100

            # folder 里找到的加分
            abs_path = os.path.abspath(path)
            abs_folder_path = os.path.abspath(folder_path) if os.path.isdir(folder_path) else None
            if abs_folder_path and abs_path.startswith(abs_folder_path + os.sep):
                score += 20

            # 有额外关键词再加分
            if must_have:
                score += 20

            try:
                mtime = os.path.getmtime(path)
            except Exception:
                mtime = 0

            candidates.append((score, mtime, path))

    if not candidates:
        return None

    # 优先：分数高 > 更新时间新 > 文件名短
    candidates.sort(key=lambda x: (-x[0], -x[1], len(os.path.basename(x[2]))))
    return candidates[0][2]


def add_records_table(doc, records, title=None, preferred_keys=None, formatters=None,
                      empty_text="No records available.", include_index=True):
    """
    将 list[dict] 渲染成表格。
    - preferred_keys: 优先显示的列顺序
    - formatters: {key: callable(value)->str}
    """
    if title:
        add_styled_heading(doc, title, level=2, center=False, size=Pt(12), bold=True)

    if records is None:
        records = []
    if isinstance(records, dict):
        records = [records]
    if not isinstance(records, list):
        records = [records]

    if not records:
        p = doc.add_paragraph()
        r = p.add_run(empty_text)
        set_run_font(r, chinese_font=CHINESE_FONT, english_font=ENGLISH_FONT)
        doc.add_paragraph()
        return

    # 收集所有字段
    all_keys = []
    seen = set()

    preferred_keys = preferred_keys or []
    for k in preferred_keys:
        if k not in seen:
            all_keys.append(k)
            seen.add(k)

    for rec in records:
        if not isinstance(rec, dict):
            continue
        for k in rec.keys():
            if k not in seen:
                all_keys.append(k)
                seen.add(k)

    cols = (1 if include_index else 0) + len(all_keys)
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    col_idx = 0
    if include_index:
        set_cell_text(hdr_cells[0], "No.", size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        col_idx = 1

    for i, key in enumerate(all_keys):
        set_cell_text(
            hdr_cells[col_idx + i],
            key,
            size=Pt(10),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER
        )

    for i, rec in enumerate(records, 1):
        row_cells = table.add_row().cells
        cidx = 0
        if include_index:
            set_cell_text(row_cells[0], str(i), size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            cidx = 1

        for j, key in enumerate(all_keys):
            value = "-"
            if isinstance(rec, dict):
                value = rec.get(key, "-")

            if formatters and key in formatters and callable(formatters[key]):
                try:
                    value_text = formatters[key](value)
                except Exception:
                    value_text = format_value_for_table(value)
            else:
                value_text = format_value_for_table(value)

            set_cell_text(
                row_cells[cidx + j],
                value_text,
                size=Pt(9),
                bold=False,
                align=WD_ALIGN_PARAGRAPH.LEFT
            )

    doc.add_paragraph()


def add_kv_table(doc, kv_rows, title=None):
    """
    添加一个 2 列键值表
    kv_rows: [(key, value), ...]
    """
    if title:
        add_styled_heading(doc, title, level=2, center=False, size=Pt(12), bold=True)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Field", size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr[1], "Value", size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for k, v in kv_rows:
        row = table.add_row().cells
        set_cell_text(row[0], str(k), size=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[1], format_value_for_table(v), size=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()


def add_sequence_table(doc, title, items, col_name="Item"):
    """
    添加一个列表表格（编号 + 内容）
    """
    add_styled_heading(doc, title, level=2, center=False, size=Pt(12), bold=True)

    if not items:
        p = doc.add_paragraph()
        r = p.add_run("None")
        set_run_font(r, chinese_font=CHINESE_FONT, english_font=ENGLISH_FONT)
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "No.", size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr[1], col_name, size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, item in enumerate(items, 1):
        row = table.add_row().cells
        set_cell_text(row[0], str(i), size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], format_value_for_table(item), size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()


def add_state_overview_section(doc, state_dicts, config, state_path=None, config_path=None, data_path=None):
    """
    顶层状态摘要：把 state 的关键信息做一个总览。
    """
    add_styled_heading(doc, "State Overview", level=1, center=False, size=Pt(14), bold=True)

    results = state_dicts.get("results", {}) or {}
    processed_files = state_dicts.get("processed_files", []) or []
    valid_sources = state_dicts.get("valid_sources", []) or []
    skipped_sources = state_dicts.get("skipped_sources", {}) or {}
    source_names = state_dicts.get("source_names", []) or []

    # 统计项
    beta_count = 0
    lsp_peak_count = 0
    jv_count = 0
    dcf_peak_count = 0
    wwz_peak_count = 0

    for src, res in results.items():
        beta_best, beta_err = extract_beta_fields(res)
        if not is_null_like(beta_best) or not is_null_like(beta_err):
            beta_count += 1

        lsp = res.get("LSP") or {}
        if isinstance(lsp, dict) and (lsp.get("periods") or []):
            lsp_peak_count += 1

        jv = res.get("Jurkevich") or {}
        if isinstance(jv, dict) and not is_null_like(jv.get("period")):
            jv_count += 1

        dcf = res.get("DCF") or {}
        if isinstance(dcf, dict) and (dcf.get("period") or []):
            dcf_peak_count += 1

        wwz = res.get("WWZ") or {}
        if isinstance(wwz, dict) and (wwz.get("result") or []):
            wwz_peak_count += 1

    summary_rows = [
        ("Report Generated At", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Data Path", data_path if data_path else "-"),
        ("State File", state_path if state_path else "-"),
        ("Config File", config_path if config_path else "-"),
        ("Processed Files Count", len(processed_files)),
        ("Source Names Count", len(source_names)),
        ("Valid Sources Count", len(valid_sources)),
        ("Skipped Sources Count", len(skipped_sources)),
        ("Successful Results Count", len(results)),
        ("Sources with Beta", beta_count),
        ("Sources with LSP Peaks", lsp_peak_count),
        ("Sources with Jurkevich Period", jv_count),
        ("Sources with DCF Peaks", dcf_peak_count),
        ("Sources with WWZ Peaks", wwz_peak_count),
    ]
    add_kv_table(doc, summary_rows, title="Top-level Summary")

    # add_sequence_table(doc, "Top-level Keys", list(state_dicts.keys()), col_name="Key")

    doc.add_paragraph()


def add_file_lists_section(doc, state_dicts):
    """
    输出 processed_files / source_names / valid_sources
    """
    add_styled_heading(doc, "File and Source Lists", level=1, center=False, size=Pt(14), bold=True)

    add_sequence_table(doc, "processed_files", state_dicts.get("processed_files", []) or [], col_name="File Name")
    add_sequence_table(doc, "valid_sources", state_dicts.get("valid_sources", []) or [], col_name="Source Name")


def add_skipped_sources_section(doc, skipped_sources):
    """
    把被跳过的源单独列出来。
    """
    if not skipped_sources:
        return

    add_styled_heading(doc, "Skipped Sources", level=1, center=False, size=Pt(14), bold=True)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "File Name", size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr[1], "Source Name", size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr[2], "Reason", size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for file_name, info in skipped_sources.items():
        row = table.add_row().cells
        set_cell_text(row[0], str(file_name), size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[1], str(info.get("source_name", "-")), size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[2], str(info.get("reason", "-")), size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()


def add_expected_periods_section(doc, state_dicts):
    """
    输出 state 中的期望/候选周期数组，按 results 的顺序对齐。
    这些数据来自：
    - lsp_expected_period
    - jv_expected_period
    - dcf_possible_period
    - wwz_possibly_period

    注意：这些数组通常与 results.keys() 的顺序对齐，不包含 skipped_sources。
    """
    ordered_sources = get_ordered_result_sources(state_dicts)
    if not ordered_sources:
        return

    lsp_expected_period = state_dicts.get("lsp_expected_period", []) or []
    jv_expected_period = state_dicts.get("jv_expected_period", []) or []
    dcf_possible_period = state_dicts.get("dcf_possible_period", []) or []
    wwz_possibly_period = state_dicts.get("wwz_possibly_period", []) or []

    add_styled_heading(doc, "State-aligned Expected / Candidate Period Arrays", level=1, center=False, size=Pt(14), bold=True)
    p = doc.add_paragraph()
    p.add_run(
        "说明：以下内容按 results 的成功源顺序对齐（即 results.keys() 的顺序），"
        "不包含 skipped_sources。"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lsp_fmt = {
        "frequency": lambda v: fmt_num(v, 6),
        "period_org": lambda v: fmt_num(v, 2),
        "period": lambda v: fmt_num(v, 2),
        "period_err": lambda v: fmt_num(v, 2),
        "power": lambda v: fmt_num(v, 4),
        "significance": lambda v: fmt_num(v, 6),
        "sigma": lambda v: fmt_num(v, 2),
    }
    dcf_fmt = {
        "period": lambda v: fmt_num(v, 2),
        "uncertainty": lambda v: fmt_num(v, 2),
        "dcf_strength": lambda v: fmt_num(v, 4),
    }

    for idx, source_name in enumerate(ordered_sources, 1):
        add_styled_heading(doc, f"{idx}. {source_name}", level=2, center=False, size=Pt(12), bold=True)

        # LSP expected
        lsp_items = lsp_expected_period[idx - 1] if idx - 1 < len(lsp_expected_period) else []
        add_records_table(
            doc,
            lsp_items,
            title="LSP Expected Periods",
            preferred_keys=["frequency", "period_org", "period", "period_err", "power", "significance", "sigma"],
            formatters=lsp_fmt,
            empty_text="No LSP expected periods."
        )

        # Jurkevich expected
        jv_item = jv_expected_period[idx - 1] if idx - 1 < len(jv_expected_period) else None
        add_kv_table(
            doc,
            [("Expected Period", fmt_num(jv_item, 2) if not is_null_like(jv_item) else "-")],
            title="Jurkevich Expected Period"
        )

        # DCF possible
        dcf_items = dcf_possible_period[idx - 1] if idx - 1 < len(dcf_possible_period) else []
        add_records_table(
            doc,
            dcf_items,
            title="DCF Possible Periods",
            preferred_keys=["period", "uncertainty", "dcf_strength"],
            formatters=dcf_fmt,
            empty_text="No DCF candidate periods."
        )

        # WWZ possible
        wwz_items = wwz_possibly_period[idx - 1] if idx - 1 < len(wwz_possibly_period) else []
        add_records_table(
            doc,
            wwz_items,
            title="WWZ Possible Periods",
            preferred_keys=None,
            formatters=None,
            empty_text="No WWZ candidate periods."
        )

    doc.add_page_break()


def add_overview_table(doc, state_dicts, config):
    """
    生成 Overview 表：
    - Source Name
    - Status
    - Effective Range
    - Beta
    - LSP / Jurkevich / DCF / WWZ 摘要
    """
    results = state_dicts.get("results", {}) or {}
    valid_sources = state_dicts.get("valid_sources", []) or []
    source_names = state_dicts.get("source_names", []) or []

    # 顺序：先 valid_sources，再 results.keys，再 source_names
    sources = []
    for s in valid_sources:
        if s in results and s not in sources:
            sources.append(s)
    for s in results.keys():
        if s not in sources:
            sources.append(s)
    for s in source_names:
        if s not in sources:
            sources.append(s)

    add_styled_heading(doc, "Overview", level=1, center=False, size=Pt(14), bold=True)

    headers = ["Source Name", "Status", "Effective Range", "Beta", "LSP", "Jurkevich", "DCF", "WWZ"]
    table = doc.add_table(rows=len(sources) + 1, cols=len(headers))
    table.style = 'Table Grid'

    for j, col_name in enumerate(headers):
        set_cell_text(
            table.cell(0, j),
            col_name,
            size=Pt(10),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER
        )

    for i, src in enumerate(sources, 1):
        res = results.get(src, {}) or {}
        date_info = resolve_source_date_info(src, res, config)

        beta_text = summarize_method_result("Beta", res)
        values = [
            src,
            res.get("status", "-"),
            date_info["effective_range"],
            beta_text,
            summarize_method_result("LSP", res.get("LSP")),
            summarize_method_result("Jurkevich", res.get("Jurkevich")),
            summarize_method_result("DCF", res.get("DCF")),
            summarize_method_result("WWZ", res.get("WWZ")),
        ]

        for j, value in enumerate(values):
            set_cell_text(
                table.cell(i, j),
                value,
                size=Pt(9),
                bold=False,
                align=WD_ALIGN_PARAGRAPH.CENTER
            )

    doc.add_paragraph()
    doc.add_page_break()


def add_source_detail_section(doc, source_name, source_result, data_path, config,
                              processed_files=None,
                              image_width=5.8, table_row_height=3.0, table_cell_width=6.5,
                              section_index=1, total_sections=1):
    """
    为单个 source 生成：
    1) 源信息
    2) 日期范围信息
    3) 统计信息
    4) 各方法结果表
    5) Method Plots
    """
    source_result = source_result or {}
    processed_files = processed_files or []

    short_name = re.sub(r'^\d+_', '', source_name)
    source_idx = extract_source_index(source_name)
    matched_file = find_processed_file_for_source(source_name, processed_files)
    date_info = resolve_source_date_info(source_name, source_result, config)
    beta_best, beta_err = extract_beta_fields(source_result)

    heading = add_styled_heading(
        doc,
        f"{section_index}. {source_name}",
        level=1,
        center=False,
        size=Pt(14),
        bold=True
    )
    add_bookmark(heading, f"src_{section_index}", bookmark_id=section_index)

    # ====== 基础信息表 ======
    metadata_rows = [
        ("Source Name", source_name),
        ("Short Name", short_name),
        ("Source Index", source_idx if source_idx is not None else "-"),
        ("Matched Processed File", matched_file if matched_file else "-"),
        ("Status", source_result.get("status", "-")),
        ("Effective Date Range", date_info["effective_range"]),
        ("Applied Date Range", date_info["applied_range"]),
        ("Config Override Date Range", date_info["override_range"]),
        ("Global Date Range", date_info["global_range"]),
    ]
    add_kv_table(doc, metadata_rows, title="Source Metadata")

    # ====== 统计信息 ======
    stats = source_result.get("stats") or {}
    stats_rows = [
        ("N_total", stats.get("N_total", "-")),
        ("N_eff", stats.get("N_eff", "-")),
        ("N_ul", stats.get("N_ul", "-")),
        ("det_ratio", fmt_num(stats.get("det_ratio"), 4)),
        ("ul_ratio", fmt_num(stats.get("ul_ratio"), 4)),
        ("P_min", fmt_num(stats.get("P_min"), 2)),
        ("Beta", fmt_num(beta_best, 4)),
        ("Beta Err", fmt_num(beta_err, 4)),
    ]
    add_kv_table(doc, stats_rows, title="Source Statistics")

    # ====== 方法摘要 ======
    method_rows = [
        ("LSP", summarize_method_result("LSP", source_result.get("LSP"))),
        ("Jurkevich", summarize_method_result("Jurkevich", source_result.get("Jurkevich"))),
        ("DCF", summarize_method_result("DCF", source_result.get("DCF"))),
        ("WWZ", summarize_method_result("WWZ", source_result.get("WWZ"))),
        ("Beta", summarize_method_result("Beta", source_result)),
    ]
    add_kv_table(doc, method_rows, title="Method Summaries")

    # ====== 详细结果表 ======
    add_styled_heading(doc, "Raw / Detailed Results", level=2, center=False, size=Pt(12), bold=True)

    # LSP
    lsp_result = source_result.get("LSP") or {}
    lsp_periods = lsp_result.get("periods") or []
    add_records_table(
        doc,
        lsp_periods,
        title="LSP Period Candidates",
        preferred_keys=["frequency", "period_org", "period", "period_err", "power", "significance", "sigma"],
        formatters={
            "frequency": lambda v: fmt_num(v, 6),
            "period_org": lambda v: fmt_num(v, 2),
            "period": lambda v: fmt_num(v, 2),
            "period_err": lambda v: fmt_num(v, 2),
            "power": lambda v: fmt_num(v, 4),
            "significance": lambda v: fmt_num(v, 6),
            "sigma": lambda v: fmt_num(v, 2),
        },
        empty_text="No LSP period candidates."
    )

    # Jurkevich
    jv_result = source_result.get("Jurkevich") or {}
    add_kv_table(
        doc,
        [
            ("period", fmt_num(jv_result.get("period"), 2)),
            ("period_err", fmt_num(jv_result.get("period_err"), 2)),
            ("boundary_list", jv_result.get("boundary_list", "-")),
        ],
        title="Jurkevich Result"
    )

    # DCF
    dcf_result = source_result.get("DCF") or {}
    dcf_periods = dcf_result.get("period") or []
    add_records_table(
        doc,
        dcf_periods,
        title="DCF Period Candidates",
        preferred_keys=["period", "uncertainty", "dcf_strength"],
        formatters={
            "period": lambda v: fmt_num(v, 2),
            "uncertainty": lambda v: fmt_num(v, 2),
            "dcf_strength": lambda v: fmt_num(v, 4),
        },
        empty_text="No DCF period candidates."
    )

    # WWZ
    wwz_result = source_result.get("WWZ") or {}
    wwz_candidates = wwz_result.get("result") or []
    add_records_table(
        doc,
        wwz_candidates,
        title="WWZ Result",
        preferred_keys=None,
        formatters=None,
        empty_text="No WWZ period candidates."
    )

    # Beta
    beta_result = source_result.get("Beta") or {}
    if not isinstance(beta_result, dict):
        beta_result = {}
    beta_rows = [
        ("beta_best", beta_result.get("beta_best", beta_best)),
        ("beta_err", beta_result.get("beta_err", beta_err)),
    ]
    add_kv_table(doc, beta_rows, title="Beta Result")

    # ====== 6×1 图片布局 ======
    add_styled_heading(doc, "Method Plots", level=2, center=False, size=Pt(12), bold=True)

    table = doc.add_table(rows=len(PLOT_SPECS), cols=1)
    table.style = 'Table Grid'
    set_no_border_table(table)

    # 只设置最小高度，让图片可以自动撑开
    for row in table.rows:
        row.height = Inches(table_row_height)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            cell.width = Inches(table_cell_width)

    # 单元格内边距设为 0
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblCellMar = parse_xml(
        r'<w:tblCellMar {}>'
        r'<w:top w:w="0" w:type="dxa"/>'
        r'<w:left w:w="0" w:type="dxa"/>'
        r'<w:bottom w:w="0" w:type="dxa"/>'
        r'<w:right w:w="0" w:type="dxa"/>'
        r'</w:tblCellMar>'.format(nsdecls('w'))
    )
    tblPr.append(tblCellMar)

    for i, spec in enumerate(PLOT_SPECS):
        cell = table.cell(i, 0)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 统一从指定 folder 中找图，避免误匹配
        image_path = find_best_image_in_folder(
            data_path,
            spec["folder"],
            source_name,
            must_have=spec.get("must_have"),
            must_not_have=spec.get("must_not_have")
        )

        if image_path is None:
            image_path = create_placeholder_image(f"No {spec['label']} Image")

        # 插图
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(image_width))

        # 方法标题
        title_paragraph = cell.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(spec["label"])
        set_run_font(
            title_run,
            chinese_font=CHINESE_FONT,
            english_font=ENGLISH_FONT,
            size=Pt(10),
            bold=True
        )

        # 方法摘要
        summary_paragraph = cell.add_paragraph()
        summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if spec["summary_method"] == "Beta":
            summary_source = source_result
        else:
            summary_source = source_result.get(spec["result_key"]) if spec["result_key"] else None

        summary_text = summarize_method_result(spec["summary_method"], summary_source)
        summary_run = summary_paragraph.add_run(summary_text)
        set_run_font(
            summary_run,
            chinese_font=CHINESE_FONT,
            english_font=ENGLISH_FONT,
            size=Pt(8),
            bold=False
        )

    doc.add_paragraph()

    if section_index < total_sections:
        doc.add_page_break()


def create_results_report_from_state(
    state_dicts,
    data_path,
    json_params,
    title="Report on the Results of the Quasi Periodic Analysis Program",
    conclusion_text=None,
    output_filename='./result_report.docx',
    image_width=6.5,
    table_row_height=3.0,
    table_cell_width=6.5,
    state_path=None,
    config_path=None
):
    """
    直接根据新版 state.json 生成 Word 报告。
    """
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    setup_footer(doc)
    setup_header(doc, "QPA-Cycler")

    # 标题
    if title:
        add_styled_heading(doc, title, level=1, center=True, size=Pt(14), bold=True)
        doc.add_paragraph()

    # 说明/结论文本
    if conclusion_text is None:
        results = state_dicts.get("results", {}) or {}
        skipped = state_dicts.get("skipped_sources", {}) or {}
        processed = state_dicts.get("processed_files", []) or []
        conclusion_text = (
            f"本报告根据 state.json 自动生成，共处理 {len(processed)} 个输入文件，"
            f"成功 {len(results)} 个源，跳过 {len(skipped)} 个源。"
            f"每个源的日期范围、统计信息、周期候选和图片结果见后续章节。"
        )

    if conclusion_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(conclusion_text.strip())
        set_run_font(r, size=Pt(10), bold=False)

    # State 总览
    add_state_overview_section(
        doc,
        state_dicts=state_dicts,
        config=json_params,
        state_path=state_path,
        config_path=config_path,
        data_path=data_path
    )

    # 参数表：把 config 展开成一行一个参数
    if json_params and isinstance(json_params, dict):
        add_parameters_table(doc, json_params, title="Analysis Parameter Settings")

    # 跳过源
    add_skipped_sources_section(doc, state_dicts.get("skipped_sources", {}))

    # 总览表
    add_overview_table(doc, state_dicts, json_params)

    # 源详情
    results = state_dicts.get("results", {}) or {}
    processed_files = state_dicts.get("processed_files", []) or []
    ordered_sources = get_ordered_result_sources(state_dicts)
    total_sections = len(ordered_sources)

    for idx, source_name in enumerate(ordered_sources, 1):
        source_result = results.get(source_name, {}) or {}
        add_source_detail_section(
            doc,
            source_name,
            source_result,
            data_path,
            json_params,
            processed_files=processed_files,
            image_width=image_width,
            table_row_height=table_row_height,
            table_cell_width=table_cell_width,
            section_index=idx,
            total_sections=total_sections
        )

    # 保存
    os.makedirs(os.path.dirname(output_filename) or ".", exist_ok=True)
    doc.save(output_filename)
    print(f"Word document saved as: {output_filename}")


# =========================================================
# 对外主函数
# =========================================================

def save2docx(
        data_path: str,
        state_dict_filename: str = 'state',
        json_params_filename: str = 'config',
        title="Report on the Results of the Quasi Periodic Analysis Program for Fermi Data",
        conclusion_text=None,
        docx_output_path='.',
        output_filename=None,
        image_width=5.8,
        table_row_height=3.0,
        table_cell_width=6.5,
        avoid_overwrite=True
):
    """
    适配新版 state.json 的真实数据导出函数。

    依赖的 state.json 结构大致为：
    {
        "processed_files": [...],
        "valid_sources": [...],
        "skipped_sources": {...},
        "results": {
            "source_name": {
                "stats": {...},
                "LSP": {
                    "periods": [...]
                } / None,
                "Jurkevich": {...} / None,
                "DCF": {...} / None,
                "WWZ": {...} / None,
                "Beta": {
                    "beta_best": ...,
                    "beta_err": ...
                },
                "applied_start_date": ...,
                "applied_end_date": ...
            }
        },
        "lsp_expected_period": [...],
        "jv_expected_period": [...],
        "dcf_possible_period": [...],
        "wwz_possibly_period": [...],
        "source_names": [...]
    }

    本版本新增：
    - 更稳健的 state/config 查找
    - Beta 从独立的 Beta 节中读取
    - 每个源的日期范围输出
    - 更完整的 state 输出

    参数说明：
    - output_filename: 自定义输出文件名（可选）
        * 如果提供完整路径，则使用该路径
        * 如果只提供文件名，则保存到 docx_output_path 目录下
        * 如果为 None，则自动生成带时间戳的文件名
    - avoid_overwrite: 是否避免覆盖已有文件（默认 True）
        * True: 如果文件已存在，自动添加序号或时间戳
        * False: 直接覆盖已有文件
    """
    candidate_dirs = [
        os.path.join(data_path, 'Running_Data'),
        data_path,
        '.'
    ]

    state_dicts, state_path = load_json_from_candidates(candidate_dirs, state_dict_filename)
    if state_dicts is None:
        raise FileNotFoundError(
            f"未找到状态文件: {os.path.join(candidate_dirs[0], _ensure_json_suffix(state_dict_filename))}"
        )

    # config 优先从 state 所在目录读，其次从 data_path，最后从当前目录读
    config_candidate_dirs = []
    if state_path:
        config_candidate_dirs.append(os.path.dirname(state_path))
    config_candidate_dirs.extend([os.path.join(data_path, 'Running_Data'), data_path, '.'])

    json_params, config_path = load_json_from_candidates(config_candidate_dirs, json_params_filename)

    # 确定输出文件路径
    os.makedirs(docx_output_path, exist_ok=True)

    if output_filename is None:
        # 未提供文件名，自动生成带时间戳的文件名
        current_date = datetime.now().strftime("%Y_%m_%d_%H_%M_%S")
        output_filename = os.path.join(docx_output_path, f"结果报告_{current_date}.docx")
    elif not os.path.isabs(output_filename):
        # 提供的是相对路径或纯文件名，拼接到输出目录
        output_filename = os.path.join(docx_output_path, output_filename)
    else:
        # 提供的是绝对路径，直接使用
        # 确保父目录存在
        os.makedirs(os.path.dirname(output_filename), exist_ok=True)

    # 防覆盖处理
    if avoid_overwrite and os.path.exists(output_filename):
        base, ext = os.path.splitext(output_filename)

        # 尝试添加序号：filename_1.docx, filename_2.docx, ...
        counter = 1
        new_filename = f"{base}_{counter}{ext}"
        while os.path.exists(new_filename):
            counter += 1
            new_filename = f"{base}_{counter}{ext}"

        output_filename = new_filename
        print(f"[save2docx] 文件已存在，自动重命名为: {os.path.basename(output_filename)}")

    create_results_report_from_state(
        state_dicts=state_dicts,
        data_path=data_path,
        json_params=json_params,
        title=title,
        conclusion_text=conclusion_text,
        output_filename=output_filename,
        image_width=image_width,
        table_row_height=table_row_height,
        table_cell_width=table_cell_width,
        state_path=state_path,
        config_path=config_path
    )

    return output_filename


# =========================================================
# main
# =========================================================

if __name__ == '__main__':
    config_map = "config"
    with open(f'{config_map}.json', encoding='utf-8') as f:
        config = json.load(f)
    output_path = r"S:\QPAwenz\FSRQ_15"
    state_name = r"state_fsrq_7.json"
    save2docx(
        data_path=output_path,
        state_dict_filename=state_name,
        json_params_filename='config',
        docx_output_path=output_path
    )